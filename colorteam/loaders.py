"""Reading proposal text out of the formats proposals actually arrive in.

Federal proposals live in Word, not Markdown. A tool that only reads .md is a
demo; one that reads .docx is usable on a Tuesday afternoon with a real draft.

Everything downstream (lint, the agents) works on plain text with line numbers,
so a loader's only job is to produce that text in a stable, reproducible order.
"""

from __future__ import annotations

from pathlib import Path

TEXT_SUFFIXES = {".md", ".markdown", ".txt", ".rst", ".text"}
DOCX_SUFFIXES = {".docx"}
SUPPORTED = TEXT_SUFFIXES | DOCX_SUFFIXES


class UnsupportedDocument(ValueError):
    pass


class MissingDependency(RuntimeError):
    pass


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _load_docx(path: Path) -> str:
    """Extract paragraphs and table cells in document order.

    Word stores tables outside the paragraph stream, so a naive paragraph walk
    silently drops every compliance table in the document — which in a proposal
    is usually where the requirements live. This walks the body XML instead, so
    the output order matches what a human reads on the page.
    """
    try:
        from docx import Document  # python-docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover - exercised by hand
        raise MissingDependency(
            "Reading .docx needs python-docx. Install it with: "
            "pip install python-docx"
        ) from exc

    document = Document(str(path))
    body = document.element.body
    lines: list[str] = []

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = Paragraph(child, document).text.strip()
            lines.append(text)
        elif tag == "tbl":
            for row in Table(child, document).rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append(" | ".join(cells))

    # Collapse runs of blank lines so line numbers stay meaningful to a reader.
    cleaned: list[str] = []
    for line in lines:
        if line or (cleaned and cleaned[-1]):
            cleaned.append(line)
    return "\n".join(cleaned).strip() + "\n"


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_comments(path: str | Path) -> list[dict]:
    """Pull reviewer comments out of a .docx, with the text each one marks.

    Review feedback arrives as Word comments in the margin, not as a separate
    document. python-docx does not expose them, so this reads the package
    directly: `word/comments.xml` holds the comment bodies, and
    `word/document.xml` holds `commentRangeStart`/`commentRangeEnd` pairs that
    bracket the text a comment is attached to.

    Returns [] for a file with no comments, and for a non-.docx file, so a
    caller can always ask.
    """
    import zipfile
    from xml.etree import ElementTree

    path = Path(path)
    if path.suffix.lower() not in DOCX_SUFFIXES or not path.exists():
        return []

    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "word/comments.xml" not in names:
                return []
            comments_xml = archive.read("word/comments.xml")
            document_xml = archive.read("word/document.xml") if "word/document.xml" in names else b""
    except (zipfile.BadZipFile, KeyError):
        return []

    anchors = _comment_anchors(document_xml) if document_xml else {}

    comments = []
    for node in ElementTree.fromstring(comments_xml).findall(f"{W}comment"):
        body = "".join(t.text or "" for t in node.iter(f"{W}t")).strip()
        if not body:
            continue
        comment_id = node.get(f"{W}id", "")
        comments.append(
            {
                "id": comment_id,
                "author": node.get(f"{W}author", "").strip(),
                "initials": node.get(f"{W}initials", "").strip(),
                "date": node.get(f"{W}date", "").strip(),
                "text": body,
                "anchor": anchors.get(comment_id, ""),
            }
        )
    return comments


def _comment_anchors(document_xml: bytes) -> dict[str, str]:
    """Map each comment id to the document text it brackets."""
    from xml.etree import ElementTree

    try:
        root = ElementTree.fromstring(document_xml)
    except ElementTree.ParseError:
        return {}

    open_ids: set[str] = set()
    collected: dict[str, list[str]] = {}
    for element in root.iter():
        tag = element.tag
        if tag == f"{W}commentRangeStart":
            ident = element.get(f"{W}id")
            if ident is not None:
                open_ids.add(ident)
                collected.setdefault(ident, [])
        elif tag == f"{W}commentRangeEnd":
            open_ids.discard(element.get(f"{W}id"))
        elif tag == f"{W}t" and open_ids and element.text:
            for ident in open_ids:
                collected[ident].append(element.text)

    return {
        ident: " ".join("".join(parts).split())[:400]
        for ident, parts in collected.items()
        if parts
    }


def format_comments(comments: list[dict]) -> str:
    """Render extracted comments as readable review feedback."""
    if not comments:
        return ""
    lines = []
    for c in comments:
        who = c["author"] or c["initials"] or "reviewer"
        when = f" · {c['date'][:10]}" if c["date"] else ""
        lines.append(f"[{c['id']}] {who}{when}")
        if c["anchor"]:
            lines.append(f'    on: "{c["anchor"]}"')
        lines.append(f"    says: {c['text']}")
        lines.append("")
    return "\n".join(lines).strip()


def load_document(path: str | Path) -> str:
    """Return the plain text of a document, chosen by file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"no such file: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _load_text(path)
    if suffix in DOCX_SUFFIXES:
        return _load_docx(path)

    raise UnsupportedDocument(
        f"cannot read {suffix or 'a file with no extension'}. "
        f"Supported: {', '.join(sorted(SUPPORTED))}"
    )
