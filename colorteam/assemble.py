"""Assembling the deliverable.

A proposal is submitted as a Word document, so a tool that stops at Markdown
stops one step short of useful. This builds the .docx: the drafted sections with
real heading styles, then the two appendices a compliant submission carries — the
compliance matrix, and a traceability table mapping every requirement to the
section that answers it.

Two choices worth naming.

**Open items are marked in the document, not silently carried.** Every
`[PROOF NEEDED]`, `[SME INPUT]`, or `[GRAPHIC]` marker is highlighted where it
sits and listed in an Open Items appendix. A marker that survives into a
submitted proposal is a fabrication risk; making it impossible to miss on the
page is cheaper than hoping someone greps for it.

**Uncovered requirements appear in the traceability table as gaps.** The table
tells the truth about what is not answered rather than omitting those rows,
because the reason to build a traceability table at all is to find the holes.
"""

from __future__ import annotations

import re
from pathlib import Path

from .matrix import Requirement

MARKER = re.compile(r"\[(PROOF NEEDED|SME INPUT|PAST PERFORMANCE|GRAPHIC)[^\]]*\]")
HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
BOLD = re.compile(r"\*\*(.+?)\*\*")


def _require_docx():
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "Building a .docx needs python-docx. Install it with: pip install python-docx"
        ) from exc


def find_markers(text: str) -> list[dict]:
    """Every open item in the draft, with the section it sits in."""
    section = "(top)"
    found = []
    for line in text.splitlines():
        heading = HEADING.match(line)
        if heading:
            section = heading.group(2)
            continue
        for match in MARKER.finditer(line):
            found.append({"section": section, "marker": match.group(0)})
    return found


def _add_runs(paragraph, text: str, highlight) -> None:
    """Write a line, bolding **spans** and highlighting open-item markers."""
    position = 0
    for match in MARKER.finditer(text):
        if match.start() > position:
            _add_plain(paragraph, text[position:match.start()])
        run = paragraph.add_run(match.group(0))
        run.bold = True
        run.font.highlight_color = highlight
        position = match.end()
    if position < len(text):
        _add_plain(paragraph, text[position:])


def _add_plain(paragraph, text: str) -> None:
    position = 0
    for match in BOLD.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def build(
    draft: str,
    rows: list[Requirement] | None = None,
    coverage: dict | None = None,
    title: str = "Technical Proposal",
    path: Path | str = "proposal.docx",
) -> Path:
    """Write the assembled proposal."""
    _require_docx()
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt

    rows = rows or []
    document = Document()

    document.add_heading(title, level=0)
    subtitle = document.add_paragraph(
        "Draft — contains open items. Not for submission."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    document.add_page_break()

    # --- body -------------------------------------------------------------
    in_code = False
    for line in draft.splitlines():
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            paragraph = document.add_paragraph(line)
            paragraph.style = document.styles["No Spacing"]
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        heading = HEADING.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2), level=level)
            continue

        if not stripped:
            continue

        if stripped.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs(paragraph, stripped[2:], WD_COLOR_INDEX.YELLOW)
            continue

        if stripped.startswith("|"):
            continue  # markdown tables are rebuilt as real tables in appendices

        paragraph = document.add_paragraph()
        _add_runs(paragraph, stripped, WD_COLOR_INDEX.YELLOW)

    # --- appendix A: compliance matrix ------------------------------------
    if rows:
        document.add_page_break()
        document.add_heading("Appendix A — Compliance Matrix", level=1)
        headers = ["ID", "Source", "Requirement", "Section", "Eval", "Status"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for cell, label in zip(table.rows[0].cells, headers):
            cell.text = label
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in rows:
            if not row.live:
                continue
            cells = table.add_row().cells
            for cell, value in zip(
                cells,
                [row.id, row.source, row.requirement, row.section,
                 row.eval_criterion, row.status],
            ):
                cell.text = value or ""

    # --- appendix B: traceability -----------------------------------------
    if coverage:
        document.add_page_break()
        document.add_heading("Appendix B — Traceability", level=1)
        document.add_paragraph(
            f"{coverage['covered']} of {coverage['requirements_live']} live "
            f"requirements are addressed "
            f"({coverage['coverage_rate'] * 100:.1f}%). Estimated length "
            f"{coverage['estimated_pages']} pages."
        )
        uncovered = coverage.get("uncovered", [])
        if uncovered:
            document.add_heading("Requirements not yet addressed", level=2)
            table = document.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            for cell, label in zip(table.rows[0].cells, ["ID", "Requirement", "Scored?"]):
                cell.text = label
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for item in uncovered:
                cells = table.add_row().cells
                cells[0].text = item["id"]
                cells[1].text = item["requirement"]
                cells[2].text = "YES" if item["eval_criterion"] else ""
        else:
            document.add_paragraph("Every live requirement is addressed.")

    # --- appendix C: open items -------------------------------------------
    markers = find_markers(draft)
    document.add_page_break()
    document.add_heading("Appendix C — Open Items", level=1)
    if markers:
        document.add_paragraph(
            f"{len(markers)} open item(s) remain in the text, highlighted in the "
            "body above. Each must be closed or removed before submission — an "
            "unclosed marker in a submitted proposal is a fabrication risk."
        )
        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        for cell, label in zip(table.rows[0].cells, ["Section", "Open item"]):
            cell.text = label
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for item in markers:
            cells = table.add_row().cells
            cells[0].text = item["section"]
            cells[1].text = item["marker"]
    else:
        document.add_paragraph("No open items remain in the text.")

    path = Path(path)
    document.save(str(path))
    return path
